from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import pandas as pd
from datetime import datetime
import os
import uuid
from typing import Dict, Any

from app.config import settings, logger

class DocumentGenerator:
    

    def __init__(self):
        self.documents_dir = settings.documents_dir
        os.makedirs(self.documents_dir, exist_ok=True)

    def generate_partner_report(self, partner_data: Dict[str, Any], analysis: Dict[str, Any]) -> Dict[str, Any]:
        
        start_time = datetime.now()

        try:
            doc = Document()

            self._setup_styles(doc)

            self._add_title(doc, f"Отчет по партнеру: {partner_data.get('trade_name')}")

            self._add_section_header(doc, "📋 Основная информация")
            self._add_partner_info(doc, partner_data)

            self._add_section_header(doc, "📊 Финансовые показатели")
            self._add_financial_info(doc, partner_data)

            self._add_section_header(doc, "🤖 Анализ искусственного интеллекта")
            self._add_ai_analysis(doc, analysis)

            self._add_section_header(doc, "📞 Контактная информация")
            self._add_contact_info(doc, partner_data)

            self._add_section_header(doc, "🏢 Адреса")
            self._add_addresses(doc, partner_data)

            if partner_data.get('financials', {}).get('turnovers'):
                self._add_section_header(doc, "📈 Исторические данные об оборотах")
                self._add_turnover_table(doc, partner_data)

            self._add_footer(doc, partner_data)

            report_uuid = str(uuid.uuid4())
            filename = f"partner_report_{partner_data.get('inn')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(self.documents_dir, filename)

            doc.save(filepath)

            file_size = os.path.getsize(filepath)

            generation_time = (datetime.now() - start_time).total_seconds() * 1000

            logger.info(f"✅ Report generated: {filename} ({file_size} bytes)")

            return {
                'filepath': filepath,
                'filename': filename,
                'file_size_bytes': file_size,
                'generation_time_ms': round(generation_time, 2),
                'report_uuid': report_uuid,
                'success': True
            }

        except Exception as e:
            logger.error(f"Error generating document: {e}")
            return {
                'error': str(e),
                'success': False
            }

    def _setup_styles(self, doc: Document):
        
        styles = doc.styles

        heading_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.size = Pt(16)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0, 0, 0)

        subheading_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        subheading_style.font.size = Pt(14)
        subheading_style.font.bold = True
        subheading_style.font.color.rgb = RGBColor(44, 62, 80)

        normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.size = Pt(11)
        normal_style.font.name = 'Calibri'

    def _add_title(self, doc: Document, title: str):
        
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run(title)
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(41, 128, 185)

        doc.add_paragraph()  # Пустая строка

    def _add_section_header(self, doc: Document, header: str):
        
        paragraph = doc.add_paragraph(header, style='CustomHeading2')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_partner_info(self, doc: Document, partner_data: Dict[str, Any]):
        
        info = [
            f"Юридическое наименование: {partner_data.get('legal_name')}",
            f"Торговое наименование: {partner_data.get('trade_name')}",
            f"ИНН: {partner_data.get('inn')}",
            f"Тип партнера: {self._translate_partner_type(partner_data.get('partner_type'))}",
            f"Категория: {partner_data.get('category')}",
            f"Основной конкурент: {partner_data.get('competitor')}",
            f"Год основания: {partner_data.get('financials', {}).get('founding_year')}",
            f"Количество сотрудников: {partner_data.get('financials', {}).get('employee_count'):,}"
        ]

        for item in info:
            p = doc.add_paragraph(item, style='CustomNormal')
            p.paragraph_format.left_indent = Inches(0.2)

    def _add_financial_info(self, doc: Document, partner_data: Dict[str, Any]):
        
        financials = partner_data.get('financials', {})

        info = [
            f"Выручка 2023 года: ${financials.get('revenue_2023', 0):,.2f}",
            f"Выручка 2022 года: ${financials.get('revenue_2022', 0):,.2f}",
            f"Прибыль 2023 года: ${financials.get('profit_2023', 0):,.2f}",
            f"Рост выручки: {self._calculate_growth(financials.get('revenue_2023', 0), financials.get('revenue_2022', 0))}%"
        ]

        for item in info:
            p = doc.add_paragraph(item, style='CustomNormal')
            p.paragraph_format.left_indent = Inches(0.2)

        ratings = partner_data.get('ratings', {})
        ratings_info = [
            f"Рейтинг: {ratings.get('rating', 0)}/5",
            f"Уровень риска: {ratings.get('risk_level', 'Unknown')}",
            f"Условия оплаты: {ratings.get('payment_terms', 'Не указано')}"
        ]

        doc.add_paragraph()  # Пустая строка

        for item in ratings_info:
            p = doc.add_paragraph(item, style='CustomNormal')
            p.paragraph_format.left_indent = Inches(0.2)

    def _add_ai_analysis(self, doc: Document, analysis: Dict[str, Any]):
        
        ai_data = analysis.get('analysis', {})

        financial = ai_data.get('financial_analysis', {})
        if financial:
            p = doc.add_paragraph("📈 Финансовый анализ:", style='CustomNormal')
            p.paragraph_format.left_indent = Inches(0.2)

            for key, value in financial.items():
                if key != 'error':
                    doc.add_paragraph(f"  • {key}: {value}", style='CustomNormal')

        risk = ai_data.get('risk_assessment', {})
        if risk:
            doc.add_paragraph()  # Пустая строка
            p = doc.add_paragraph("⚠️ Оценка рисков:", style='CustomNormal')
            p.paragraph_format.left_indent = Inches(0.2)

            doc.add_paragraph(f"  • Уровень: {risk.get('level', 'Unknown')}", style='CustomNormal')

            if risk.get('factors'):
                doc.add_paragraph("  • Факторы риска:", style='CustomNormal')
                for factor in risk.get('factors', []):
                    doc.add_paragraph(f"    - {factor}", style='CustomNormal')

            if risk.get('recommendations'):
                doc.add_paragraph("  • Рекомендации:", style='CustomNormal')
                for rec in risk.get('recommendations', []):
                    doc.add_paragraph(f"    - {rec}", style='CustomNormal')

        potential = ai_data.get('partnership_potential', {})
        if potential:
            doc.add_paragraph()  # Пустая строка
            p = doc.add_paragraph("🤝 Потенциал партнерства:", style='CustomNormal')
            p.paragraph_format.left_indent = Inches(0.2)

            doc.add_paragraph(f"  • Оценка: {potential.get('score', 0)}/10", style='CustomNormal')

            if potential.get('opportunities'):
                doc.add_paragraph("  • Возможности:", style='CustomNormal')
                for opp in potential.get('opportunities', []):
                    doc.add_paragraph(f"    - {opp}", style='CustomNormal')

            if potential.get('threats'):
                doc.add_paragraph("  • Угрозы:", style='CustomNormal')
                for threat in potential.get('threats', []):
                    doc.add_paragraph(f"    - {threat}", style='CustomNormal')

        recommendations = ai_data.get('strategic_recommendations', [])
        if recommendations:
            doc.add_paragraph()  # Пустая строка
            p = doc.add_paragraph("💡 Стратегические рекомендации:", style='CustomNormal')
            p.paragraph_format.left_indent = Inches(0.2)

            for rec in recommendations:
                doc.add_paragraph(f"  • {rec}", style='CustomNormal')

        summary = ai_data.get('summary', '')
        if summary:
            doc.add_paragraph()  # Пустая строка
            p = doc.add_paragraph("🎯 Итоговый вывод:", style='CustomNormal')
            p.paragraph_format.left_indent = Inches(0.2)
            doc.add_paragraph(summary, style='CustomNormal')

    def _add_contact_info(self, doc: Document, partner_data: Dict[str, Any]):
        
        contacts = partner_data.get('contacts', {})

        info = [
            f"Генеральный директор (CEO): {contacts.get('ceo', 'Не указан')}",
            f"Финансовый директор (CFO): {contacts.get('cfo', 'Не указан')}",
            f"Email для связи: {contacts.get('email', 'Не указан')}",
            f"Контактный телефон: {contacts.get('phone', 'Не указан')}"
        ]

        for item in info:
            p = doc.add_paragraph(item, style='CustomNormal')
            p.paragraph_format.left_indent = Inches(0.2)

        if partner_data.get('website'):
            p = doc.add_paragraph(f"Веб-сайт: {partner_data.get('website')}", style='CustomNormal')
            p.paragraph_format.left_indent = Inches(0.2)

    def _add_addresses(self, doc: Document, partner_data: Dict[str, Any]):
        
        addresses = partner_data.get('addresses', [])

        if addresses:
            for address in addresses:
                p = doc.add_paragraph(f"• {address}", style='CustomNormal')
                p.paragraph_format.left_indent = Inches(0.2)
        else:
            doc.add_paragraph("Адреса не указаны", style='CustomNormal')

    def _add_turnover_table(self, doc: Document, partner_data: Dict[str, Any]):
        
        turnovers = partner_data.get('financials', {}).get('turnovers', [])

        if not turnovers:
            return

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Shading'

        headers = ['Год', 'Квартал', 'Выручка ($)', 'Прибыль ($)', 'Средний чек ($)']
        hdr_cells = table.rows[0].cells

        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True

        for turnover in turnovers:
            row_cells = table.add_row().cells
            row_cells[0].text = str(turnover.get('year', ''))
            row_cells[1].text = str(turnover.get('quarter', '')) if turnover.get('quarter') else 'Год'
            row_cells[2].text = f"${turnover.get('revenue', 0):,.2f}"
            row_cells[3].text = f"${turnover.get('profit', 0):,.2f}" if turnover.get('profit') else 'N/A'
            row_cells[4].text = f"${turnover.get('average_transaction', 0):,.2f}" if turnover.get('average_transaction') else 'N/A'

    def _add_footer(self, doc: Document, partner_data: Dict[str, Any]):
        
        doc.add_page_break()

        footer = doc.sections[0].footer
        paragraph = footer.paragraphs[0]

        current_time = datetime.now().strftime("%d.%m.%Y %H:%M")

        footer_text = (
            f"Отчет сгенерирован автоматически: {current_time} | "
            f"Партнер: {partner_data.get('trade_name')} | "
            f"ИНН: {partner_data.get('inn')}"
        )

        paragraph.text = footer_text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.runs[0].font.size = Pt(9)
        paragraph.runs[0].font.italic = True
        paragraph.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    def _translate_partner_type(self, partner_type: str) -> str:
        
        translations = {
            'strategic': 'Стратегический',
            'current': 'Текущий',
            'potential': 'Потенциальный',
            'blocked': 'Заблокированный',
            'vip': 'VIP'
        }
        return translations.get(partner_type, partner_type)

    def _calculate_growth(self, current: float, previous: float) -> str:
        
        if not previous or previous == 0:
            return "N/A"

        growth = ((current - previous) / previous) * 100
        return f"{growth:+.1f}"

document_generator = DocumentGenerator()
